import streamlit as st
import os, re
import io
from glob import glob
import pandas as pd
from docx import Document
st.set_page_config(layout='wide')

# ----- Functions
@st.cache_data
def read_table(table_file):
    if table_file.name.endswith('.csv'):
        df = pd.read_csv(table_file)
    else:
        df = pd.read_excel(table_file)

    return df

# @st.cache_resource
def read_doc(doc_file):
    doc = Document(doc_file)
    return doc

# 템플릿 문서의 {{}} 내부 공백 제거
def remove_space(text):
    return re.sub(r"\{\{\s*|\s*\}\}", lambda match: match.group().replace(" ", ""), text)

# 템플릿 문서 내의 변수명 추출
def get_key_names(text):
    var_names = re.findall(r"\{\{([^\{\}]+)\}\}", text)
    return list(set([x.strip() for x in var_names]))

# 변수 -> 텍스트 대치
def replace_text(text, df):
    for k, v in df[['key', 'value']].values:
        text = text.replace(f'{{{{{k}}}}}', str(v))

    return text

    


st.header('Template-Based Report Generator')

# ----- Sample files
st.markdown('#### 샘플 데이터 다운로드')

with open('./sample/Audit_Report_Variables_Example_Values.csv') as f:
    st.download_button('Download CSV', f, 'Audit_Report_Variables_Example_Values.csv')

with open('sample/Audit_Report_Template.docx', 'rb') as f:
    st.download_button('Download DOCX', f, 'Audit_Report_Template.docx')


# ----- Data upload
st.divider()
st.markdown('#### 데이터 업로드')
up_col1, up_col2 = st.columns(2)
with up_col1:
    table_upload = st.file_uploader(label='변수 테이블 업로드', type=['csv', 'xlsx'], key='table_file_uploader')

with up_col2:
    doc_upload = st.file_uploader(label='문서 템플릿 업로드', type=['docx'], key='doc_file_uploader')

if not table_upload or not doc_upload:
    st.stop()

# ---- Read data
st.divider()
st.markdown('#### 데이터 확인')
df = read_table(table_upload)
doc = read_doc(doc_upload)
# doc_ = read_doc(doc_upload)

doc_full_text = '\n'.join([x.text for x in doc.paragraphs])

table_col, doc_col = st.columns((2,4))
with table_col:
    st.write('**변수 테이블**')
    st.dataframe(df, height=500)

with doc_col:
    st.write('**템플릿 문서**')
    st.text_area(label='doc text', value=doc_full_text[:1000], height=500, label_visibility='collapsed', disabled=True)


doc_full_text = remove_space(doc_full_text) # 전처리: 템플릿 문서의 {{}} 내부 공백 제거

doc_key_names = get_key_names(doc_full_text) # 템플릿 문서에 있는 변수명 리스트
csv_key_names = df['key'].unique().tolist() # 사용자 입력 변수 테이블에 있는 변수명 리스트

inter_keys = set(doc_key_names).intersection(set(csv_key_names)) # 매칭된 변수
if len(inter_keys) == len(doc_key_names) == len(csv_key_names):
    st.info('✅ 사용자 입력 변수 테이블과 템플릿 문서의 변수가 완전히 일치합니다.')

else:
    warning_message = "⚠️ 사용자 입력 변수 테이블과 템플릿 문서의 변수가 일치하지 않습니다. 매칭 성공된 변수면 대치됩니다.\n"
    warning_message += f'- 매칭 성공된 변수: {list(inter_keys)}\n'

    doc_left_keys = list(set(doc_key_names) - inter_keys) # 템플릿 문서에는 존재하지만, 사용자 입력 테이블에 없는 변수명
    csv_left_keys = list(set(csv_key_names) - inter_keys) # 사용자 입력 테이블에는 존재하지만, 템플릿 문서에는 없는 변수명
    if len(doc_left_keys):
        warning_message += f'- 사용자 입력 테이블에 다음 변수가 없습니다: {doc_left_keys}\n'

    if len(csv_left_keys):
        warning_message += f'- 템플릿 문서 내에 다음 변수가 없습니다: {csv_left_keys}'

    st.warning(warning_message)


# ----- Replace text
st.divider()
st.markdown('### 문서 생성')
with st.spinner('In progress...'):
    # 절 단위 탐색 및 대치
    for p in doc.paragraphs:
        p.text = replace_text(text=remove_space(p.text), df=df)

    # 테이블 탐색 및 대치
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = replace_text(text=remove_space(para.text), df=df)

# save_path = doc_path.replace('.docx', '_Result.docx')
# doc.save(save_path)

save_docx_name = 'Result_' + doc_upload.name
bio = io.BytesIO()
doc.save(bio)
st.download_button(label='Download Result DOCX', data=bio.getvalue(), file_name=save_docx_name, mime='docx', type='primary')